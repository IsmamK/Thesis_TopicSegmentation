import os
import pandas as pd
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# ======= CONFIG =======
SUMMARY_CSV = "4_Results/Baseline/summary_metrics.csv"
PLOTS_DIR = "4_Results/Baseline/outputs"
REPORT_PATH = "4_Results/Baseline/Baseline_Report.docx"

# ======= INIT DOC =======
doc = Document()
doc.add_heading("Baseline Text Segmentation Report", level=0)

doc.add_paragraph("This report summarizes the baseline approach for automatic topic segmentation "
                  "of lecture videos using text-based sentence embeddings and cosine similarity drop detection. "
                  "The approach was applied to five videos using Whisper transcripts.")

doc.add_paragraph("A Sentence-BERT model ('all-MiniLM-L6-v2') was used to encode sentence segments from Whisper-generated transcripts. "
                  "Then, cosine similarity was calculated between consecutive embeddings. A smoothing filter was applied, "
                  "and significant drops in similarity were used to infer topic boundaries via peak detection.")

doc.add_paragraph("Ground truth boundaries were extracted from YouTube chapter metadata, and evaluation was performed "
                  "using Precision, Recall, and F1 Score (with a 5-second tolerance window).")

# ======= LOAD METRICS =======
doc.add_heading("Evaluation Summary", level=1)

if not os.path.exists(SUMMARY_CSV):
    doc.add_paragraph("❌ Could not find summary_metrics.csv. Please run evaluation first.")
else:
    df = pd.read_csv(SUMMARY_CSV)

    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = col

    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)

    doc.add_paragraph("\nInterpretation:")

    for idx, row in df.iterrows():
        doc.add_paragraph(f"• {row['video']}: Predicted {row['predicted_boundaries']} boundaries "
                          f"vs {row['ground_truth_boundaries']} ground truth. "
                          f"F1 Score = {row['f1']}.", style="List Bullet")

# ======= ADD PLOTS =======
doc.add_heading("Cosine Similarity Drop Plots", level=1)

for filename in sorted(os.listdir(PLOTS_DIR)):
    if filename.endswith(".png") and "_plot" in filename:
        video_title = filename.replace("_plot.png", "")
        doc.add_heading(video_title, level=2)

        image_path = os.path.join(PLOTS_DIR, filename)
        try:
            doc.add_picture(image_path, width=Inches(6.0))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        except Exception as e:
            doc.add_paragraph(f"⚠️ Could not add image: {filename} ({e})")

# ======= SAVE =======
doc.add_page_break()
doc.save(REPORT_PATH)
print(f"✅ Report saved to: {REPORT_PATH}")
